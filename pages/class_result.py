"""
Description: 课程处理结果页面

-*- Encoding: UTF-8 -*-
@File     ：class_result.py
@Time     ：2025/2/9
"""
import time
import streamlit as st
from docx import Document
from utils.utils import navigation, check_login_state, get_response_connie
from utils.styles import apply_global_styles
import re
from datetime import datetime
from langchain.memory import ConversationBufferMemory
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import io
import zipfile
from configs.log_config import *


def extract_course_info(filename):
    """从文件名中提取日期时间和课程名称"""
    try:
        # 匹配日期时间和课程名称
        # 例如: "2025-02-05 08_10新A_导读.docx"
        pattern = r"(\d{4}-\d{2}-\d{2})\s*(\d{2}[_:]\d{2})([^_\.]*)_?"
        match = re.search(pattern, filename)

        if match:
            date_str, time_str, course_name = match.groups()
            # 处理时间中的下划线
            time_str = time_str.replace('_', ':')
            # 提取课程名称（移除可能的前后空格）
            course_name = course_name.strip()
            return date_str, time_str, course_name

        # 如果无法匹配完整格式，尝试至少提取日期和时间
        date_time_pattern = r"(\d{4}-\d{2}-\d{2})\s*(\d{2}[_:]\d{2})"
        match = re.search(date_time_pattern, filename)
        if match:
            date_str, time_str = match.groups()
            time_str = time_str.replace('_', ':')
            # 使用默认课程名称
            return date_str, time_str, "英语课程"

        # 如果都无法匹配，使用当前时间和默认课程名称
        now = datetime.now()
        return now.strftime('%Y-%m-%d'), now.strftime('%H:%M'), "英语课程"

    except Exception as e:
        st.error(f"提取课程信息时发生错误：{str(e)}")
        now = datetime.now()
        return now.strftime('%Y-%m-%d'), now.strftime('%H:%M'), "英语课程"


def extract_chapter_overview(doc_content):
    """从文档内容中提取章节速览部分"""
    try:
        # 替换所有可能的换行符组合为统一的换行符
        doc_content = doc_content.replace('\r\n', '\n').replace('\r', '\n')

        # 使用更精确的正则表达式
        pattern = r"章节速览[\s\n]*(.*?)(?=\s*(?:要点回顾|问答回顾))"
        match = re.search(pattern, doc_content, re.DOTALL | re.MULTILINE | re.IGNORECASE)

        if match:
            # 提取匹配的内容并清理
            chapter_content = match.group(1).strip()
            return chapter_content

        st.warning("未找到标准格式的章节速览")
        return None
    except Exception as e:
        st.error(f"提取章节速览时发生错误：{str(e)}")
        return None


def generate_summary(date_str, time_str, course_name, chapter_overview):
    """调用API生成课程总结"""
    try:
        # 构建prompt
        prompt = f"""你现在是一个拥有三十年教学经验的初中英语老师，你刚刚完成一节英语课的授课，以下是记录的课堂授记录：

时间：{date_str} {time_str}
名称：{course_name}

课章节速览：
{chapter_overview}

请帮我进行润色，丰富内容，按照指定格式撰写一篇专业的课堂总结。总结包括两部分主要内容：课程概述/对学生课下学习的建议。请按照1234等要点对课堂概述进行提炼。主要总结课堂上讲授了什么知识，其他无关紧要的不要总结。

总结不要分太多级。指定格式：
授课时间：{date_str}
课程名称：{course_name}
课程总结：{{content of 课程总结}}
课后建议：{{content of 课后建议}}"""

        # 创建对话记忆对象
        memory = ConversationBufferMemory()
        max_retries = 3
        retry_delay = 3  # 重试延迟时间（秒）

        # 创建一个可重用的警告消息占位符
        warning_placeholder = st.empty()

        for attempt in range(max_retries):
            try:
                # 调用API获取响应
                response = get_response_connie(prompt, memory)
                if response and isinstance(response, str) and len(response.strip()) > 0:
                    # 验证响应是否包含必要的内容
                    required_keywords = ["课程总结", "课后建议"]
                    if all(keyword in response for keyword in required_keywords):
                        # 清除警告消息
                        warning_placeholder.empty()
                        return response

                # 如果响应不符合要求，等待后重试
                if attempt < max_retries - 1:
                    warning_placeholder.warning(f"API响应不符合要求，正在进行第{attempt + 2}次尝试...")
                    time.sleep(retry_delay)

            except Exception as e:
                if attempt < max_retries - 1:
                    warning_placeholder.warning(
                        "DeepSeek API调用失败，该问题由近期DeepSeek所遭受的大规模网络攻击造成，与您无关。"
                        f"正在进行第{attempt + 2}次尝试...请耐心等待..."
                    )
                    time.sleep(retry_delay)
                else:
                    warning_placeholder.empty()
                    raise Exception(f"多次尝试后API调用仍然失败：{str(e)}")

        warning_placeholder.empty()
        raise Exception("无法获取有效的API响应")

    except Exception as e:
        st.error(f"生成总结时发生错误：{str(e)}")
        logger.error(f"生成总结时发生错误：{str(e)}")
        return None


def create_summary_document(summary_text, original_filename):
    """创建总结文档"""
    try:
        doc = Document()

        # 清理文本中的*符号
        cleaned_text = summary_text.replace('*', '')

        # 设置文档标题
        title = doc.add_paragraph("课程总结")
        title_format = title.runs[0]
        title_format.font.name = 'Microsoft YaHei'
        title_format.font.size = Pt(16)
        title_format.font.bold = True
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # 添加总结内容
        content = doc.add_paragraph(cleaned_text)
        content_format = content.runs[0]
        content_format.font.name = 'Microsoft YaHei'
        content_format.font.size = Pt(12)

        # 保存到内存中
        doc_binary = io.BytesIO()
        doc.save(doc_binary)
        doc_binary.seek(0)

        return doc_binary
    except Exception as e:
        st.error(f"创建文档时发生错误：{str(e)}")
        return None


@st.dialog("确认返回")
def confirm_return_dialog():
    """确认返回的对话框"""
    st.write("您确定要返回吗？返回后处理好的文件将被清除，如果需要请先下载。")

    # 添加确认复选框
    confirm_checkbox = st.checkbox(
        "我确认不需要下载或已完成下载",
        key="confirm_return_checkbox"
    )

    col1, col2 = st.columns(2)
    with col1:
        if st.button(
                "确认返回",
                type="primary",
                use_container_width=True,
                disabled=not confirm_checkbox  # 如果没有勾选确认框，禁用返回按钮
        ):
            # 清除session state中的所有相关数据
            if 'files_to_process' in st.session_state:
                del st.session_state['files_to_process']
            if 'generated_docs' in st.session_state:
                del st.session_state['generated_docs']
            if 'processing_complete' in st.session_state:
                del st.session_state['processing_complete']
            # 返回课程总结页面
            st.switch_page("pages/zongjie.py")

    with col2:
        if st.button("取消", use_container_width=True):
            st.rerun()


def process_documents(uploaded_files):
    """处理上传的文档"""
    try:
        # 用于收集所有生成的文档
        generated_docs = []
        failed_files = []
        total_files = len(uploaded_files)
        processed_count = 0

        # 创建等待提示
        wait_message = st.empty()
        last_message_time = time.time()
        message_index = 0
        waiting_messages = [
            "正在分析课程内容，请稍候...",
            "正在调用AI助手生成总结，请耐心等待...",
            "正在整理课程要点，请稍等...",
            "AI正在进行内容润色，马上就好..."
        ]

        # 立即显示第一条消息
        wait_message.info(
            f"{waiting_messages[0]}\n"
            f"当前处理进度：{processed_count + 1}/{total_files}"
        )

        with st.spinner("正在处理文档..."):
            for uploaded_file in uploaded_files:
                try:
                    # 更新等待消息
                    current_time = time.time()
                    if current_time - last_message_time > 3:  # 改为每3秒更新一次
                        wait_message.info(
                            f"{waiting_messages[message_index % len(waiting_messages)]}\n"
                            f"当前处理进度：{processed_count + 1}/{total_files}"
                        )
                        last_message_time = current_time
                        message_index += 1

                    # 提取文档信息
                    doc = Document(uploaded_file)
                    doc_content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                    date_str, time_str, course_name = extract_course_info(uploaded_file.name)
                    chapter_overview = extract_chapter_overview(doc_content)

                    if not chapter_overview:
                        failed_files.append((uploaded_file.name, "未找到章节速览部分"))
                        processed_count += 1
                        continue

                    # 生成总结
                    summary = generate_summary(date_str, time_str, course_name, chapter_overview)

                    if not summary:
                        failed_files.append((uploaded_file.name, "生成总结失败"))
                        processed_count += 1
                        continue

                    # 创建文档
                    doc_binary = create_summary_document(summary, uploaded_file.name)

                    if not doc_binary:
                        failed_files.append((uploaded_file.name, "创建文档失败"))
                        processed_count += 1
                        continue

                    # 保存生成的文档信息
                    output_filename = f"{date_str}_{course_name}_课程总结.docx"
                    generated_docs.append((output_filename, doc_binary))

                except Exception as e:
                    failed_files.append((uploaded_file.name, str(e)))
                finally:
                    processed_count += 1

        # 清除等待提示
        wait_message.empty()

        # 存储处理结果到session state
        st.session_state['generated_docs'] = generated_docs
        st.session_state['failed_files'] = failed_files

        # 标记处理已完成
        st.session_state['processing_complete'] = True

        # 使用 rerun 来刷新页面显示结果
        st.rerun()

    except Exception as e:
        st.error(f"处理文档时发生错误：{str(e)}")
        if 'files_to_process' in st.session_state:
            del st.session_state['files_to_process']


def class_result():
    st.set_page_config(page_title='ATM-Cleaning', page_icon='images/favicon.png')
    apply_global_styles()
    login_state, role = check_login_state()

    if not login_state:
        error = st.error("您还没有登录！3秒后跳转至登录页面...", icon="⚠️")
        time.sleep(1)
        error.empty()
        error = st.error("您还没有登录！2秒后跳转至登录页面...", icon="⚠️")
        time.sleep(1)
        error.empty()
        st.error("您还没有登录！1秒后跳转至登录页面...", icon="⚠️")
        time.sleep(1)
        st.switch_page("pages/login_page.py")
    else:
        # 检查是否是特定用户
        current_user = st.session_state.get("logged_in_username")
        if current_user != "connie":
            error = st.error("您没有权限访问此页面！3秒后跳转至登录页面...", icon="⚠️")
            time.sleep(1)
            error.empty()
            error = st.error("您没有权限访问此页面！2秒后跳转至登录页面...", icon="⚠️")
            time.sleep(1)
            error.empty()
            st.error("您没有权限访问此页面！1秒后跳转至登录页面...", icon="⚠️")
            time.sleep(1)
            st.switch_page("pages/login_page.py")
        else:
            navigation()
            st.title("📚自动化课程总结")
            st.divider()

            # 检查是否有要处理的文件
            if 'files_to_process' not in st.session_state:
                st.warning("没有需要处理的文件，请返回上传页面。", icon="⚠️")
                if st.button("返回上传页面", type="primary"):
                    st.switch_page("pages/zongjie.py")
                return

            # 如果处理已完成，显示结果
            if st.session_state.get('processing_complete', False):
                if 'generated_docs' in st.session_state:
                    st.success(f"共处理完成 {len(st.session_state['generated_docs'])} 个文件！")

                    # 显示处理完成的文件列表
                    st.write("已完成处理的文件：")
                    for filename, _ in st.session_state['generated_docs']:
                        st.text(filename)

                    st.divider()  # 添加分隔线

                    col1, col2 = st.columns(2)
                    with col1:
                        # 显示下载按钮
                        if len(st.session_state['generated_docs']) == 1:
                            filename, doc_binary = st.session_state['generated_docs'][0]
                            st.download_button(
                                label="下载课程总结",
                                data=doc_binary,
                                file_name=filename,
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                type="primary",
                                use_container_width=True
                            )
                        else:
                            # 创建zip文件
                            zip_buffer = io.BytesIO()
                            with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                                for filename, doc_binary in st.session_state['generated_docs']:
                                    zip_file.writestr(filename, doc_binary.getvalue())

                            st.download_button(
                                label="下载所有课程总结",
                                data=zip_buffer.getvalue(),
                                file_name="课程总结集.zip",
                                mime="application/zip",
                                type="primary",
                                use_container_width=True
                            )

                    with col2:
                        if st.button("返回上传页面", use_container_width=True):
                            confirm_return_dialog()

                # 显示失败的文件列表
                if 'failed_files' in st.session_state and st.session_state['failed_files']:
                    st.error("以下文件处理失败：")
                    for filename, reason in st.session_state['failed_files']:
                        st.write(f"- {filename}：{reason}")
            else:
                # 开始处理文档
                process_documents(st.session_state['files_to_process'])


if __name__ == '__main__':
    class_result()
