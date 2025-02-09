"""
Description: 课程总结页面

-*- Encoding: UTF-8 -*-
@File     ：zongjie.py
@Time     ：2025/2/9
"""
import time
import streamlit as st
from docx import Document
from utils.utils import navigation, check_login_state, get_response_connie
from utils.styles import apply_global_styles
import re
from datetime import datetime
from langchain.memory import ConversationBufferMemory
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import io
from configs.log_config import *


def extract_course_info(filename):
    """从文件名中提取日期时间和课程名称"""
    try:
        # 匹配日期时间和课程名称
        # 例如: "2025-02-05 08_10新A_导读.docx"
        pattern = r"(\d{4}-\d{2}-\d{2})\s*(\d{2}[_:]\d{2})([^_\.]*)_?"
        match = re.search(pattern, filename)

        if match:
            date_str, time_str, course_name = match.groups()
            # 处理时间中的下划线
            time_str = time_str.replace('_', ':')
            # 提取课程名称（移除可能的前后空格）
            course_name = course_name.strip()
            return date_str, time_str, course_name

        # 如果无法匹配完整格式，尝试至少提取日期和时间
        date_time_pattern = r"(\d{4}-\d{2}-\d{2})\s*(\d{2}[_:]\d{2})"
        match = re.search(date_time_pattern, filename)
        if match:
            date_str, time_str = match.groups()
            time_str = time_str.replace('_', ':')
            # 使用默认课程名称
            return date_str, time_str, "英语课程"

        # 如果都无法匹配，使用当前时间和默认课程名称
        now = datetime.now()
        return now.strftime('%Y-%m-%d'), now.strftime('%H:%M'), "英语课程"

    except Exception as e:
        st.error(f"提取课程信息时发生错误：{str(e)}")
        now = datetime.now()
        return now.strftime('%Y-%m-%d'), now.strftime('%H:%M'), "英语课程"


def extract_chapter_overview(doc_content):
    """从文档内容中提取章节速览部分"""
    try:
        # 替换所有可能的换行符组合为统一的换行符
        doc_content = doc_content.replace('\r\n', '\n').replace('\r', '\n')

        # 使用更宽松的正则表达式查找章节速览部分
        pattern = r"章节速览[\s\n]*?((?:[\s\S]*?(?=-\s*\d{2}:\d{2})|[\s\S]*?(?=问答回顾)|[\s\S]*?$))"
        match = re.search(pattern, doc_content, re.DOTALL)

        if match:
            # 提取匹配的内容并清理
            chapter_content = match.group(1).strip()

            # 使用新的正则表达式提取所有时间段内容
            time_sections = re.findall(r'-\s*\d{2}:\d{2}.*?(?=(?:-\s*\d{2}:\d{2}|\s*$))', chapter_content, re.DOTALL)

            if time_sections:
                # 合并所有时间段内容
                final_content = '\n'.join(section.strip() for section in time_sections)
                return final_content

            # 如果没有找到时间段，但内容不为空，返回清理后的内容
            if chapter_content:
                return chapter_content

        st.warning("未找到标准格式的章节速览，尝试提取时间段内容...")

        # 如果上述方法都失败，直接尝试提取所有时间段内容
        time_sections = re.findall(r'-\s*\d{2}:\d{2}.*?(?=(?:-\s*\d{2}:\d{2}|\s*$))', doc_content, re.DOTALL)
        if time_sections:
            return '\n'.join(section.strip() for section in time_sections)

        return None
    except Exception as e:
        st.error(f"提取章节速览时发生错误：{str(e)}")
        return None


def generate_summary(date_str, time_str, course_name, chapter_overview):
    """调用API生成课程总结"""
    try:
        # 构建prompt
        prompt = f"""你现在是一个拥有三十年教学经验的初中英语老师，你刚刚完成一节英语课的授课，以下是记录的课堂授记录：

时间：{date_str} {time_str}
名称：{course_name}

课章节速览：
{chapter_overview}

请帮我进行润色，丰富内容，按照指定格式撰写一篇专业的课堂总结。总结包括两部分主要内容：课程概述/对学生课下学习的建议。请按照1234等要点对课堂概述进行提炼。主要总结课堂上讲授了什么知识，其他无关紧要的不要总结。

总结不要分太多级。指定格式：
授课时间：{date_str}
课程名称：{course_name}
课程总结：{{content of 课程总结}}
课后建议：{{content of 课后建议}}"""

        # 创建对话记忆对象
        memory = ConversationBufferMemory()
        max_retries = 3
        retry_delay = 3  # 重试延迟时间（秒）

        for attempt in range(max_retries):
            try:
                # 调用API获取响应
                response = get_response_connie(prompt, memory)
                if response and isinstance(response, str) and len(response.strip()) > 0:
                    # 验证响应是否包含必要的内容
                    required_keywords = ["课程总结", "课后建议"]
                    if all(keyword in response for keyword in required_keywords):
                        return response

                # 如果响应不符合要求，等待后重试
                if attempt < max_retries - 1:
                    st.warning(f"API响应不符合要求，正在进行第{attempt + 2}次尝试...")
                    time.sleep(retry_delay)

            except Exception as e:
                if attempt < max_retries - 1:
                    st.warning(f"API调用失败，正在进行第{attempt + 2}次尝试...")
                    time.sleep(retry_delay)
                else:
                    raise Exception(f"多次尝试后API调用仍然失败：{str(e)}")

        raise Exception("无法获取有效的API响应")

    except Exception as e:
        st.error(f"生成总结时发生错误：{str(e)}")
        logger.error(f"生成总结时发生错误：{str(e)}")
        return None


def create_summary_document(summary_text, original_filename):
    """创建总结文档"""
    try:
        doc = Document()

        # 清理文本中的*符号
        cleaned_text = summary_text.replace('*', '')

        # 设置文档标题
        title = doc.add_paragraph("课程总结")
        title_format = title.runs[0]
        title_format.font.name = 'Microsoft YaHei'
        title_format.font.size = Pt(16)
        title_format.font.bold = True
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # 添加总结内容
        content = doc.add_paragraph(cleaned_text)
        content_format = content.runs[0]
        content_format.font.name = 'Microsoft YaHei'
        content_format.font.size = Pt(12)

        # 保存到内存中
        doc_binary = io.BytesIO()
        doc.save(doc_binary)
        doc_binary.seek(0)

        return doc_binary
    except Exception as e:
        st.error(f"创建文档时发生错误：{str(e)}")
        return None


def process_documents(uploaded_files):
    """处理上传的文档"""
    try:
        # 用于收集所有生成的文档
        generated_docs = []
        failed_files = []
        total_files = len(uploaded_files)
        processed_count = 0

        for uploaded_file in uploaded_files:
            # 创建信息提示的占位符
            info_placeholder = st.empty()
            info_placeholder.info(f"正在处理：{uploaded_file.name} ({processed_count + 1}/{total_files})")

            # 提取文档信息
            doc = Document(uploaded_file)
            doc_content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            date_str, time_str, course_name = extract_course_info(uploaded_file.name)
            chapter_overview = extract_chapter_overview(doc_content)

            if not chapter_overview:
                info_placeholder.empty()
                failed_files.append((uploaded_file.name, "未找到章节速览部分"))
                processed_count += 1
                continue

            # 生成总结
            summary = generate_summary(date_str, time_str, course_name, chapter_overview)

            if not summary:
                info_placeholder.empty()
                failed_files.append((uploaded_file.name, "生成总结失败"))
                processed_count += 1
                continue

            # 创建文档
            doc_binary = create_summary_document(summary, uploaded_file.name)

            if not doc_binary:
                info_placeholder.empty()
                failed_files.append((uploaded_file.name, "创建文档失败"))
                processed_count += 1
                continue

            # 保存生成的文档信息
            output_filename = f"{date_str}_{course_name}_课程总结.docx"
            generated_docs.append((output_filename, doc_binary))

            # 清除处理中提示
            info_placeholder.empty()
            processed_count += 1

        # 处理完成后，显示结果
        if generated_docs:
            st.success(f"共处理完成 {len(generated_docs)} 个文件！")

            # 如果只有一个文档
            if len(generated_docs) == 1:
                filename, doc_binary = generated_docs[0]
                st.download_button(
                    label="下载课程总结",
                    data=doc_binary,
                    file_name=filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    type="primary"
                )
            # 如果有多个文档，创建zip文件
            else:
                import io
                import zipfile

                # 创建内存中的zip文件
                zip_buffer = io.BytesIO()
                with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                    for filename, doc_binary in generated_docs:
                        zip_file.writestr(filename, doc_binary.getvalue())

                # 提供zip文件下载
                st.download_button(
                    label="下载所有课程总结",
                    data=zip_buffer.getvalue(),
                    file_name="课程总结集.zip",
                    mime="application/zip",
                    type="primary"
                )

        # 显示失败的文件列表
        if failed_files:
            st.error("以下文件处理失败：")
            for filename, reason in failed_files:
                st.write(f"- {filename}：{reason}")

    except Exception as e:
        st.error(f"处理文档时发生错误：{str(e)}")


def course_summary():
    st.set_page_config(page_title='ATM-Cleaning', page_icon='images/favicon.png')
    apply_global_styles()
    login_state, role = check_login_state()

    if not login_state:
        error = st.error("您还没有登录！3秒后跳转至登录页面...", icon="⚠️")
        time.sleep(1)
        error.empty()
        error = st.error("您还没有登录！2秒后跳转至登录页面...", icon="⚠️")
        time.sleep(1)
        error.empty()
        st.error("您还没有登录！1秒后跳转至登录页面...", icon="⚠️")
        time.sleep(1)
        st.switch_page("pages/login_page.py")
    else:
        # 检查是否是特定用户
        current_user = st.session_state.get("logged_in_username")
        if current_user != "connie":
            error = st.error("您没有权限访问此页面！3秒后跳转至登录页面...", icon="⚠️")
            time.sleep(1)
            error.empty()
            error = st.error("您没有权限访问此页面！2秒后跳转至登录页面...", icon="⚠️")
            time.sleep(1)
            error.empty()
            st.error("您没有权限访问此页面！1秒后跳转至登录页面...", icon="⚠️")
            time.sleep(1)
            st.switch_page("pages/login_page.py")
        else:
            navigation()
            st.title("📚课程总结")
            st.divider()

            # 添加说明信息
            st.info("请上传您需要处理的记录文件，该文件应从通义听悟中直接导出，不要对导出文件进行任何修改。", icon="ℹ️")

            # 初始化session state来存储已处理的文件名
            if 'processed_files' not in st.session_state:
                st.session_state.processed_files = set()

            # 文件上传部分
            uploaded_files = st.file_uploader(
                "选择要处理的Word文档",
                type=['docx'],
                accept_multiple_files=True,
                key='file_uploader'
            )

            if uploaded_files:
                # 检查文件名是否重复
                current_files = set()
                valid_files = []

                for file in uploaded_files:
                    if file.name not in current_files and file.name not in st.session_state.processed_files:
                        current_files.add(file.name)
                        valid_files.append(file)

                if valid_files and st.button("开始处理", type="primary"):
                    with st.spinner("正在处理文档..."):
                        process_documents(valid_files)
                        # 添加处理过的文件名到session state
                        st.session_state.processed_files.update(current_files)


if __name__ == '__main__':
    course_summary()