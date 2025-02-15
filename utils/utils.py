"""
Description: 项目中用到的工具函数
    
-*- Encoding: UTF-8 -*-
@File     ：utils.py.py
@Author   ：King Songtao
@Time     ：2024/12/26 下午8:27
@Contact  ：king.songtao@gmail.com
"""
from datetime import date

import streamlit as st
import time
import re
from docx.shared import Pt
from langchain.memory import ConversationBufferMemory
from configs.log_config import *


def stream_res(res):
    """前端制作流式输出效果"""
    for char in res:
        yield char
        time.sleep(0.02)


def check_login_state():
    """获取或存储登录状态"""
    try:
        if "login_state" not in st.session_state:
            st.session_state["login_state"] = False
            st.session_state["role"] = None
            return False, st.session_state["role"]
        elif st.session_state["login_state"]:
            # 检查用户是否在活跃会话中
            username = st.session_state.get("logged_in_username")
            if username and username in get_active_sessions():
                return True, st.session_state["role"]
            else:
                # 如果不在活跃会话中，清除登录状态
                st.session_state["login_state"] = False
                st.session_state["role"] = None
                return False, None
        else:
            st.session_state["login_state"] = False
            st.session_state["role"] = None
            return False, st.session_state["role"]
    except Exception as e:
        logger.error(f"检测登录状态时发生错误！错误信息：{e}")
        st.error("发生未知错误！即将跳转到首页...")
        st.switch_page("pages/login_page.py")

    # """通过cookies管理登陆状态"""
    # if cookies.get("is_logged_in") == "1":
    #     st.session_state['is_logged_in'] = True
    #     st.session_state['role'] = cookies.get("role", "user")
    #     return True, cookies.get("role")
    # else:
    #     st.session_state['is_logged_in'] = False
    #     return False, cookies.get("role")


def set_login_state(is_logged_in, role, name):
    st.session_state["login_state"] = True if is_logged_in else False
    st.session_state["role"] = role
    st.session_state["name"] = name

    # cookies["is_logged_in"] = "1" if is_logged_in else "0"
    # cookies["role"] = role
    # cookies["name"] = name
    # cookies.save()  # 保存 Cookie


def log_out():
    """注销登录"""
    username = st.session_state.get("logged_in_username")
    if username:
        remove_active_session(username)

    st.session_state["login_state"] = False
    st.session_state["role"] = None
    st.switch_page("app.py")


def validate_address(address):
    """
    验证地址是否只包含英文字符、数字和常见标点符号
    """
    # 增加 / 到允许的字符中
    pattern = r'^[a-zA-Z0-9\s\.\,\-\#\/]+$'

    if not address:
        return False, "地址不能为空"

    if not re.match(pattern, address):
        return False, "地址只能包含英文字符、数字和符号(.,- #/)。请检查是否含有中文逗号！"

    return True, ""


def extract_date_from_html(html_content):
    """
    提取 HTML 内容中的日期
    假设日期格式为 "16th Dec. 2024"（可以根据需要调整正则表达式）
    """
    # 用正则表达式匹配日期
    date_pattern = r"\d{1,2}[a-zA-Z]{2}\s[A-Za-z]+\.\s\d{4}"
    match = re.search(date_pattern, html_content)

    if match:
        return match.group(0)  # 返回匹配到的日期
    return None


def generate_receipt(doc, data_dict):
    """
        替换文档中的占位符并统一字体
        """
    # 替换段落中的文本
    for paragraph in doc.paragraphs:
        for key, value in data_dict.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, str(value))

        # 统一段落字体和大小
        for run in paragraph.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(10)

    # 替换表格中的文本
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in data_dict.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, str(value))

                # 统一单元格内段落的字体和大小
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Arial'
                        run.font.size = Pt(10)

    return doc


def formate_date(input_date):
    """
    自定义日期格式化
    将 2024-12-11 格式化为 11th Dec. 2024
    """
    day = input_date.day
    month_abbr = input_date.strftime('%b')
    year = input_date.year

    # 添加日期后缀
    if 10 <= day % 100 <= 20:
        suffix = 'th'
    else:
        suffix = {1: 'st', 2: 'nd', 3: 'rd'}.get(day % 10, 'th')

    return f"{day}{suffix} {month_abbr}. {year}"


def formate_acc_info(data):
    formatted_orders = []
    for i in range(len(data['账户编号'])):
        formatted_order = f"{data['登录账号'][i]} | {data['用户名'][i]}"
        formatted_orders.append(formatted_order)
    return formatted_orders


@st.dialog("退出登录！")
def confirm_logout():
    st.write("您正在退出登录，确认要继续吗?")
    col1, col2 = st.columns([0.5, 0.5])
    with col1:
        if st.button("确认", key="confirm_logout", type="primary", use_container_width=True):
            log_out()
    with col2:
        if st.button("取消", key="cancel_logout", type="secondary", use_container_width=True):
            st.rerun()


def navigation():
    # 导航模块
    st.sidebar.image("images/logo.png")

    if st.sidebar.button("➕创建收据", key="open_receipt_button", use_container_width=True, type="primary"):
        clear_form_state()
        st.switch_page("pages/receipt_page.py")

    # 自动化报价
    if st.sidebar.button("➕创建工单", key="auto_quote_button", use_container_width=True, type="primary"):
        st.switch_page("pages/new_work_order_v2.py")

    # 工单管理
    if st.sidebar.button("🔍工单管理", key="order_management", use_container_width=True, type="primary"):
        st.switch_page("pages/orders_statistics.py")

    if st.sidebar.button("👩‍👩‍👧‍👦月度结算", key="staff_management_button", use_container_width=True, type="primary"):
        st.switch_page("pages/monthly_review.py")

    # 人员管理模块
    if st.sidebar.button("👥人员管理", key="user_management_button", use_container_width=True, type="primary"):
        st.switch_page("pages/staff_acc.py")

    # 仅当登录用户为connie时显示课程总结按钮
    if st.session_state.get("logged_in_username") == "connie":
        if st.sidebar.button("📚课程总结", key="course_summary_button", use_container_width=True, type="primary"):
            st.switch_page("pages/zongjie.py")

    st.sidebar.divider()

    # 控制台
    if st.sidebar.button("📊回控制台", key="admin_page_button", use_container_width=True):
        st.switch_page("pages/orders_statistics.py")

    # 个人设置
    if st.sidebar.button("⚙️系统设置", key="system_setting_button", use_container_width=True):
        st.switch_page("pages/system_setting.py")

    # 退出登录模块
    if st.sidebar.button("🛏️退出登录", key="logout_button", use_container_width=True):
        confirm_logout()

    st.sidebar.write("Copy Right © 2025 ATM Cleaning Management")


def clear_form_state():
    """清空表单状态"""
    empty_state = {
        "selected_template": "手动版（手动选择excluded中的内容）",
        "address": "",
        "selected_date": date.today(),
        "amount": 0.0,
        "basic_service": [],
        "electrical": [],
        "rooms": [],
        "other": [],
        "custom_notes": [],
        "custom_notes_enabled": False,
        "excluded_enabled": False,
        "custom_excluded_enabled": False,  # 新增字段
        "manual_excluded_selection": [],
        "custom_excluded_items": [],
        "output_doc": None,
        "receipt_file_name": "",
        "ready_doc": None
    }

    # 清除session state
    for key in ['previous_form_data', 'included_items', 'excluded_items']:
        if key in st.session_state:
            del st.session_state[key]

    st.session_state['previous_form_data'] = empty_state
    st.switch_page("pages/receipt_page.py")


@st.dialog("请选择操作")
def confirm_back():
    st.write("您要修改该收据还是重开新收据？选择重开新收据，已经录入的内容将被清空！")
    col1, col2 = st.columns(2)
    with col1:
        if st.button("修改收据", use_container_width=True, type="primary"):
            if 'receipt_data' in st.session_state:
                current_data = st.session_state['receipt_data'].copy()

                # 确保保存完整的自定义项目内容
                if 'included_items' in st.session_state:
                    current_data['custom_notes'] = st.session_state['included_items']
                if 'excluded_items' in st.session_state:
                    current_data['custom_excluded_items'] = st.session_state['excluded_items']

                st.session_state['previous_form_data'] = current_data
            st.switch_page("pages/receipt_page.py")
    with col2:
        if st.button("重开收据", use_container_width=True, type="secondary"):
            clear_form_state()


def get_response(prompt, memory):
    """
    获取AI响应的函数

    Args:
        prompt (str): 用户输入的提示词
        memory (ConversationBufferMemory): 对话记忆对象

    Returns:
        str: AI的响应文本
    """
    from langchain.chat_models import ChatOpenAI
    from langchain.chains import ConversationChain

    # 从streamlit secrets中获取API配置
    chat_model = ChatOpenAI(
        model="deepseek-chat",
        openai_api_key=st.secrets["api_keys"]["openai_api_key"],
        openai_api_base=st.secrets["api_keys"]["openai_api_base"],
        temperature=0.7
    )

    # 构建对话链
    chain = ConversationChain(
        llm=chat_model,
        memory=memory,
        verbose=True
    )

    try:
        if not isinstance(prompt, str):
            prompt = str(prompt)

        response = chain.run(prompt)
        return response

    except Exception as e:
        logger.error(f"Error in get_response: {str(e)}")
        return f"抱歉，生成回复时出现错误：{str(e)}"


def get_response_connie(prompt, memory):
    """
    获取AI响应的函数

    Args:
        prompt (str): 用户输入的提示词
        memory (ConversationBufferMemory): 对话记忆对象

    Returns:
        str: AI的响应文本

    Raises:
        Exception: 当无法获取有效响应时抛出异常
    """
    from langchain.chat_models import ChatOpenAI
    from langchain.chains import ConversationChain
    import time

    # 重试参数
    max_retries = 3
    retry_delay = 5  # 增加重试间隔
    last_error = None

    for attempt in range(max_retries):
        try:
            # 确保prompt是字符串
            if not isinstance(prompt, str):
                prompt = str(prompt)

            # 从streamlit secrets中获取API配置
            chat_model = ChatOpenAI(
                model="deepseek-chat",
                openai_api_key=st.secrets["api_keys"]["openai_api_key"],
                openai_api_base=st.secrets["api_keys"]["openai_api_base"],
                temperature=0.7,
                request_timeout=600  # 设置超时时间为10分钟
            )

            # 构建对话链
            chain = ConversationChain(
                llm=chat_model,
                memory=memory,
                verbose=True
            )

            # 获取响应
            response = chain.run(prompt)

            # 验证响应
            if not response or not isinstance(response, str) or len(response.strip()) == 0:
                raise ValueError("API返回了空响应或无效响应")

            # 验证响应中是否包含必要的关键词
            required_keywords = ["课程总结", "课后建议"]
            if not all(keyword in response for keyword in required_keywords):
                raise ValueError("API响应缺少必要的内容结构")

            return response

        except Exception as e:
            last_error = str(e)
            logger.error(f"API调用失败 (尝试 {attempt + 1}/{max_retries}): {last_error}")

            if attempt < max_retries - 1:
                time.sleep(retry_delay)
                continue

            # 所有重试都失败后，抛出异常
            raise Exception(f"多次尝试后API调用仍然失败: {last_error}")


def get_active_sessions():
    """获取所有活跃的会话"""
    if 'active_sessions' not in st.session_state:
        st.session_state.active_sessions = {}
    return st.session_state.active_sessions


def add_active_session(username):
    """添加活跃会话"""
    sessions = get_active_sessions()
    sessions[username] = True


def remove_active_session(username):
    """移除活跃会话"""
    sessions = get_active_sessions()
    if username in sessions:
        del sessions[username]


def get_theme_color():
    """
    从 .streamlit/config.toml 读取主题色
    Returns:
        str: 主题色（十六进制颜色代码）
    """
    config_path = ".streamlit/config.toml"
    default_color = "#FF4B4B"

    try:
        if os.path.exists(config_path):
            with open(config_path, "r", encoding='utf-8') as f:
                config = toml.load(f)
                return config.get("theme", {}).get("primaryColor", default_color)
        return default_color
    except Exception:
        return default_color
